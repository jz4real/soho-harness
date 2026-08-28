#!/usr/bin/env python3
"""Generate a formal Chinese meeting proposal from a small JSON payload.

The template supplies page settings and document package defaults. This script
replaces the template body as a whole so that every placeholder is removed and
the output has a predictable structure.
"""
import argparse
import json
import platform
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_CJK_FONT = 'Microsoft YaHei' if platform.system() == 'Windows' else 'Hiragino Sans GB'


def set_run_font(run, east_asia, size, bold=False):
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(document, text, *, font=DEFAULT_CJK_FONT, size=16,
                  bold=False, first_indent=True, align=None, before=0, after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(28)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_indent:
        fmt.first_line_indent = Pt(32)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return paragraph


def require_string(data, key):
    value = data.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f'Missing required string: {key}')
    return value.strip()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--template', required=True)
    parser.add_argument('--data', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()

    with open(args.data, encoding='utf-8') as fh:
        data = json.load(fh)
    title = require_string(data, 'title')
    meeting_type = require_string(data, 'meeting_type')
    submit_unit = require_string(data, 'submit_unit')
    submit_date = require_string(data, 'submit_date')
    decision = require_string(data, 'decision')
    sections = data.get('sections')
    if not isinstance(sections, list) or not sections:
        raise ValueError('sections must be a non-empty array')
    attachments = data.get('attachments', [])
    if not isinstance(attachments, list):
        raise ValueError('attachments must be an array')

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(args.template, output)
    document = Document(output)
    document._element.body.clear_content()  # retain template section settings

    add_paragraph(document, title, font=DEFAULT_CJK_FONT, size=22, bold=True,
                  first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    for section in sections:
        heading = require_string(section, 'heading')
        paragraphs = section.get('paragraphs')
        if not isinstance(paragraphs, list) or not paragraphs:
            raise ValueError(f'section {heading} needs paragraphs')
        add_paragraph(document, heading, font=DEFAULT_CJK_FONT, size=16, bold=True,
                      first_indent=True, before=6)
        for item in paragraphs:
            if not isinstance(item, str) or not item.strip():
                raise ValueError(f'section {heading} contains empty paragraph')
            add_paragraph(document, item.strip())

    add_paragraph(document, f'提请{meeting_type}审议/决策事项：{decision}',
                  font=DEFAULT_CJK_FONT, size=16, first_indent=False, before=8)
    add_paragraph(document, '附件：', font=DEFAULT_CJK_FONT, size=16,
                  first_indent=False, before=8)
    if attachments:
        for index, attachment in enumerate(attachments, start=1):
            if not isinstance(attachment, str) or not attachment.strip():
                raise ValueError('attachments must contain non-empty strings')
            add_paragraph(document, f'{index}. {attachment.strip()}', first_indent=False)
    else:
        add_paragraph(document, '无。', first_indent=False)

    add_paragraph(document, submit_unit, first_indent=False,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, before=16)
    add_paragraph(document, submit_date, first_indent=False,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    document.save(output)
    print(output)


if __name__ == '__main__':
    main()
