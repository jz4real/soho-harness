#!/usr/bin/env python3
"""Read one DOCX into a bounded, model-friendly text form for this skill."""
import argparse
import sys
from docx import Document

MAX_CHARS = 18000


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('input_docx')
    args = parser.parse_args()
    document = Document(args.input_docx)
    blocks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f'[表格 {table_index}]')
        for row in table.rows:
            values = [' '.join(cell.text.split()) for cell in row.cells]
            blocks.append(' | '.join(values))
    output = '\n'.join(blocks)
    if len(output) > MAX_CHARS:
        output = output[:MAX_CHARS] + '\n[材料超过 18000 字符，已截断；请向用户索取更聚焦的材料。]'
    sys.stdout.write(output + '\n')


if __name__ == '__main__':
    main()
